from pdf2docx import Converter
from docx import Document
import requests
import json
import time

def make_docx_from_pdf(file_path: str):
    save_path = file_path.replace(".pdf", ".docx")
    
    # Create a PdfDocument object
    cv = Converter(file_path)
    # Convert a PDF file into Docx 
    cv.convert(save_path) # all pages by default
    cv.close()



def make_translation(file_path: str):
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    print(len(paragraphs))
    
    for i, para in enumerate(paragraphs):
        try:
            res = requests.post("http://localhost:8000/translate/", data=json.dumps({'text': para}))
            
            doc.add_paragraph(res.json()['translated'])
            time.sleep(0.1)
            print("Success "+str(i))
        except:
            print("Error "+str(i))
    doc.save("translated.docx")
    print("Document translation is completed.")
    pass
    